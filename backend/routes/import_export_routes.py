"""
Data Import/Export Routes - CSV and Excel support
"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query
from fastapi.responses import StreamingResponse
from typing import Optional
import csv
import io
import json
from datetime import datetime, timezone
from bson import ObjectId
from openpyxl import Workbook, load_workbook
from config.database import get_tenant_db


EXPORTABLE_COLLECTIONS = {
    "products": {
        "fields": ["name_ar", "name_en", "barcode", "article_code", "retail_price", "wholesale_price", "purchase_price", "quantity", "min_stock", "category", "family_id", "unit", "tax_rate"],
        "label_ar": "المنتجات",
        "label_fr": "Products"
    },
    "customers": {
        "fields": ["name", "phone", "email", "address", "city", "wilaya", "notes", "family_id"],
        "label_ar": "الزبائن",
        "label_fr": "Customers"
    },
    "suppliers": {
        "fields": ["name", "phone", "email", "address", "city", "company", "tax_id", "notes"],
        "label_ar": "الموردين",
        "label_fr": "Suppliers"
    },
    "employees": {
        "fields": ["name", "phone", "email", "position", "salary", "hire_date", "notes"],
        "label_ar": "الموظفين",
        "label_fr": "Employees"
    },
    "sales": {
        "fields": ["invoice_number", "customer_name", "total", "discount", "payment_method", "payment_type", "status", "created_at", "note"],
        "label_ar": "المبيعات",
        "label_fr": "Sales"
    },
    "purchases": {
        "fields": ["invoice_number", "supplier_name", "total", "discount", "payment_method", "status", "created_at", "note"],
        "label_ar": "المشتريات",
        "label_fr": "Purchases"
    },
    "expenses": {
        "fields": ["title", "amount", "category", "payment_method", "date", "notes", "recurring"],
        "label_ar": "المصاريف",
        "label_fr": "Expenses"
    },
    "debts": {
        "fields": ["customer_name", "amount", "remaining", "type", "status", "due_date", "created_at", "notes"],
        "label_ar": "الديون",
        "label_fr": "Debts"
    }
}


def create_import_export_routes(db, get_current_user) -> dict:
    router = APIRouter(prefix="/data", tags=["data-import-export"])


    def _resolve_db(user, tenant_id):
        """Super admin can target a specific tenant's database via tenant_id."""
        if tenant_id and user.get("role") == "super_admin":
            return get_tenant_db(tenant_id)
        return db

    @router.get("/collections")
    async def get_exportable_collections(tenant_id: Optional[str] = None, user: dict = Depends(get_current_user)):
        """Get list of collections available for import/export"""
        tdb = _resolve_db(user, tenant_id)
        result = []
        for key, info in EXPORTABLE_COLLECTIONS.items():
            count = await tdb[key].count_documents({})
            result.append({
                "key": key,
                "label_ar": info["label_ar"],
                "label_fr": info["label_fr"],
                "fields": info["fields"],
                "count": count
            })
        return result

    @router.get("/export/{collection}")
    async def export_data(
        collection: str,
        format: str = Query("csv", enum=["csv", "xlsx", "txt", "pdf", "docx"]),
        tenant_id: Optional[str] = None,
        user: dict = Depends(get_current_user)
    ):
        """Export collection data as CSV or Excel"""
        if collection not in EXPORTABLE_COLLECTIONS:
            raise HTTPException(status_code=400, detail=f"Collection '{collection}' not exportable")

        tdb = _resolve_db(user, tenant_id)
        fields = EXPORTABLE_COLLECTIONS[collection]["fields"]
        cursor = tdb[collection].find({}).sort("created_at", -1)
        docs = await cursor.to_list(length=50000)

        if format == "csv":
            output = io.StringIO()
            writer = csv.DictWriter(output, fieldnames=["id"] + fields, extrasaction='ignore')
            writer.writeheader()
            for doc in docs:
                row = {"id": str(doc.get("_id", ""))}
                for f in fields:
                    val = doc.get(f, "")
                    if isinstance(val, ObjectId):
                        val = str(val)
                    elif isinstance(val, datetime):
                        val = val.isoformat()
                    elif isinstance(val, (dict, list)):
                        val = json.dumps(val, ensure_ascii=False)
                    row[f] = val or ""
                writer.writerow(row)

            output.seek(0)
            filename = f"{collection}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.csv"
            return StreamingResponse(
                iter([output.getvalue()]),
                media_type="text/csv",
                headers={"Content-Disposition": f'attachment; filename="{filename}"'}
            )

        elif format == "xlsx":  # xlsx
            wb = Workbook()
            ws = wb.active
            ws.title = collection
            headers = ["id"] + fields
            ws.append(headers)

            for doc in docs:
                row = [str(doc.get("_id", ""))]
                for f in fields:
                    val = doc.get(f, "")
                    if isinstance(val, ObjectId):
                        val = str(val)
                    elif isinstance(val, datetime):
                        val = val.isoformat()
                    elif isinstance(val, (dict, list)):
                        val = json.dumps(val, ensure_ascii=False)
                    row.append(val or "")
                ws.append(row)

            # Auto-width columns
            for col in ws.columns:
                max_length = max(len(str(cell.value or "")) for cell in col)
                ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)

            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            filename = f"{collection}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.xlsx"
            return StreamingResponse(
                output,
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f'attachment; filename="{filename}"'}
            )

        elif format == "txt":
            # Tab-separated UTF-8 text (opens cleanly in Excel/notepad)
            def _txt_row(doc):
                vals = [str(doc.get("_id", ""))]
                for f in fields:
                    val = doc.get(f, "")
                    if isinstance(val, ObjectId):
                        val = str(val)
                    elif isinstance(val, datetime):
                        val = val.isoformat()
                    elif isinstance(val, (dict, list)):
                        val = json.dumps(val, ensure_ascii=False)
                    vals.append(str(val).replace("\t", " ").replace("\n", " ") if val is not None else "")
                return vals

            lines = ["\t".join(["id"] + fields)]
            for doc in docs:
                lines.append("\t".join(_txt_row(doc)))
            payload = "\ufeff" + "\n".join(lines)  # BOM for Excel UTF-8 detection
            filename = f"{collection}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.txt"
            return StreamingResponse(
                iter([payload]),
                media_type="text/plain; charset=utf-8",
                headers={"Content-Disposition": f'attachment; filename="{filename}"'}
            )

        elif format == "pdf":
            # Arabic-aware PDF table (Noto Naskh Arabic + reshaping)
            import os as _os
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import arabic_reshaper
            from bidi.algorithm import get_display

            font_path = _os.path.join(_os.path.dirname(__file__), "..", "assets", "NotoNaskhArabic-Regular.ttf")
            pdfmetrics.registerFont(TTFont("Arabic", font_path))

            def _ar(v):
                txt = str(v)
                try:
                    return get_display(arabic_reshaper.reshape(txt))
                except Exception:
                    return txt

            def _pdf_row(doc):
                vals = [str(doc.get("_id", ""))[:12]]
                for f in fields:
                    val = doc.get(f, "")
                    if isinstance(val, ObjectId):
                        val = str(val)
                    elif isinstance(val, datetime):
                        val = val.isoformat()
                    elif isinstance(val, (dict, list)):
                        val = json.dumps(val, ensure_ascii=False)
                    vals.append(str(val) if val is not None else "")
                return [_ar(v) for v in vals]

            data = [[_ar(h) for h in ["id"] + fields]] + [_pdf_row(d) for d in docs]
            buf = io.BytesIO()
            pdf_doc = SimpleDocTemplate(buf, pagesize=landscape(A4),
                                        rightMargin=18, leftMargin=18, topMargin=18, bottomMargin=18)
            tbl = Table(data, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("FONT", (0, 0), (-1, -1), "Arabic", 7.5),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e3a5f")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]))
            pdf_doc.build([tbl])
            buf.seek(0)
            filename = f"{collection}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.pdf"
            return StreamingResponse(
                buf,
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{filename}"'}
            )

        else:  # docx
            import docx as _docx
            d = _docx.Document()
            table = d.add_table(rows=1, cols=len(["id"] + fields))
            table.style = "Table Grid"
            for i, h in enumerate(["id"] + fields):
                table.rows[0].cells[i].text = str(h)
            for doc in docs:
                cells = table.add_row().cells
                vals = [str(doc.get("_id", ""))]
                for f in fields:
                    val = doc.get(f, "")
                    if isinstance(val, ObjectId):
                        val = str(val)
                    elif isinstance(val, datetime):
                        val = val.isoformat()
                    elif isinstance(val, (dict, list)):
                        val = json.dumps(val, ensure_ascii=False)
                    vals.append(str(val) if val is not None else "")
                for i, v in enumerate(vals):
                    cells[i].text = v
            buf = io.BytesIO()
            d.save(buf)
            buf.seek(0)
            filename = f"{collection}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.docx"
            return StreamingResponse(
                buf,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{filename}"'}
            )

    @router.post("/import/{collection}")
    async def import_data(
        collection: str,
        file: UploadFile = File(...),
        mode: str = Query("append", enum=["append", "replace"]),
        tenant_id: Optional[str] = Query(None),
        user: dict = Depends(get_current_user)
    ):
        """Import data from CSV or Excel file"""
        if collection not in EXPORTABLE_COLLECTIONS:
            raise HTTPException(status_code=400, detail=f"Collection '{collection}' not importable")

        if user.get("role") not in ["admin", "super_admin"]:
            raise HTTPException(status_code=403, detail="Admin only")

        tdb = _resolve_db(user, tenant_id)
        fields = EXPORTABLE_COLLECTIONS[collection]["fields"]
        content = await file.read()
        filename = file.filename.lower()
        records = []

        try:
            if filename.endswith('.csv') or filename.endswith('.txt'):
                text = content.decode('utf-8-sig')
                if filename.endswith('.txt'):
                    first_line = text.splitlines()[0] if text.splitlines() else ''
                    delim = '\t' if '\t' in first_line else ','
                    reader = csv.DictReader(io.StringIO(text), delimiter=delim)
                else:
                    reader = csv.DictReader(io.StringIO(text))
                for row in reader:
                    record = {}
                    for f in fields:
                        if f in row and row[f]:
                            val = row[f]
                            # Type conversion
                            if f in ["retail_price", "wholesale_price", "purchase_price", "amount", "total", "discount", "salary", "remaining", "tax_rate", "quantity", "min_stock"]:
                                try:
                                    val = float(val)
                                except (ValueError, TypeError):
                                    val = 0
                            record[f] = val
                    record["created_at"] = datetime.now(timezone.utc).isoformat()
                    record["updated_at"] = datetime.now(timezone.utc).isoformat()
                    records.append(record)

            elif filename.endswith('.xlsx') or filename.endswith('.xls'):
                wb = load_workbook(io.BytesIO(content), read_only=True)
                ws = wb.active
                rows = list(ws.iter_rows(values_only=True))
                if not rows:
                    raise HTTPException(status_code=400, detail="Empty file")

                headers_row = [str(h).strip() if h else "" for h in rows[0]]
                for row in rows[1:]:
                    record = {}
                    for i, val in enumerate(row):
                        if i < len(headers_row) and headers_row[i] in fields:
                            field = headers_row[i]
                            if val is None:
                                val = ""
                            if field in ["retail_price", "wholesale_price", "purchase_price", "amount", "total", "discount", "salary", "remaining", "tax_rate", "quantity", "min_stock"]:
                                try:
                                    val = float(val) if val else 0
                                except (ValueError, TypeError):
                                    val = 0
                            record[field] = val
                    record["created_at"] = datetime.now(timezone.utc).isoformat()
                    record["updated_at"] = datetime.now(timezone.utc).isoformat()
                    records.append(record)
                wb.close()
            elif filename.endswith('.docx'):
                import docx as _docx
                d = _docx.Document(io.BytesIO(content))
                rows_data = []
                if d.tables:
                    for r in d.tables[0].rows:
                        rows_data.append([c.text.strip() for c in r.cells])
                else:
                    for para in d.paragraphs:
                        if para.text.strip():
                            rows_data.append([v.strip() for v in para.text.split('\t')])
                if not rows_data:
                    raise HTTPException(status_code=400, detail="Empty file")
                headers_row = rows_data[0]
                for row in rows_data[1:]:
                    record = {}
                    for i, val in enumerate(row):
                        if i < len(headers_row) and headers_row[i] in fields:
                            field = headers_row[i]
                            if field in ["retail_price", "wholesale_price", "purchase_price", "amount", "total", "discount", "salary", "remaining", "tax_rate", "quantity", "min_stock"]:
                                try:
                                    val = float(val) if val else 0
                                except (ValueError, TypeError):
                                    val = 0
                            record[field] = val
                    record["created_at"] = datetime.now(timezone.utc).isoformat()
                    record["updated_at"] = datetime.now(timezone.utc).isoformat()
                    records.append(record)

            elif filename.endswith('.pdf'):
                raise HTTPException(
                    status_code=400,
                    detail="استيراد PDF غير مدعوم — ملفات PDF للعرض والطباعة فقط. للاستيراد استخدم Excel أو Word أو TXT أو CSV.",
                )
                import re as _re
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(content))
                lines = []
                for page in reader.pages:
                    for ln in (page.extract_text() or '').splitlines():
                        if ln.strip():
                            lines.append(ln.strip())
                hi = None
                for i, ln in enumerate(lines):
                    if sum(1 for f in fields if f in ln) >= 2:
                        hi = i
                        break
                if hi is None:
                    raise HTTPException(status_code=400, detail="Could not find a data table in the PDF")

                def _split_ln(ln):
                    if '\t' in ln:
                        return [v.strip() for v in ln.split('\t')]
                    if '|' in ln:
                        return [v.strip() for v in ln.split('|')]
                    return [v.strip() for v in _re.split(r'\s{2,}', ln)]

                headers_row = _split_ln(lines[hi])
                for ln in lines[hi + 1:]:
                    vals = _split_ln(ln)
                    record = {}
                    for i, val in enumerate(vals):
                        if i < len(headers_row) and headers_row[i] in fields:
                            field = headers_row[i]
                            if field in ["retail_price", "wholesale_price", "purchase_price", "amount", "total", "discount", "salary", "remaining", "tax_rate", "quantity", "min_stock"]:
                                try:
                                    val = float(val) if val else 0
                                except (ValueError, TypeError):
                                    val = 0
                            record[field] = val
                    if record:
                        record["created_at"] = datetime.now(timezone.utc).isoformat()
                        record["updated_at"] = datetime.now(timezone.utc).isoformat()
                        records.append(record)

            else:
                raise HTTPException(status_code=400, detail="Unsupported file format. Use .csv, .xlsx, .txt, .docx or .pdf")

        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading file: {str(e)}")

        if not records:
            raise HTTPException(status_code=400, detail="No valid records found")

        # Log the import
        import_log = {
            "collection": collection,
            "filename": file.filename,
            "mode": mode,
            "records_count": len(records),
            "user_id": str(user.get("_id", "")),
            "user_name": user.get("name", ""),
            "created_at": datetime.now(timezone.utc).isoformat()
        }

        if mode == "replace":
            await tdb[collection].delete_many({})

        inserted = 0
        skipped_dups = 0
        if records:
            from pymongo.errors import BulkWriteError
            try:
                # unordered: valid rows still import when some hit unique indexes
                result = await tdb[collection].insert_many(records, ordered=False)
                inserted = len(result.inserted_ids)
            except BulkWriteError as bwe:
                inserted = bwe.details.get("nInserted", 0)
                skipped_dups = len(bwe.details.get("writeErrors", []))
            import_log["inserted_count"] = inserted
            import_log["skipped_duplicates"] = skipped_dups

        await tdb["import_logs"].insert_one(import_log)

        return {
            "success": True,
            "message": f"Imported {inserted} records to {collection}"
                       + (f" ({skipped_dups} duplicates skipped)" if skipped_dups else ""),
            "records_imported": inserted,
            "skipped_duplicates": skipped_dups,
            "mode": mode
        }

    @router.get("/import-history")
    async def get_import_history(tenant_id: Optional[str] = None, user: dict = Depends(get_current_user)):
        """Get import history"""
        tdb = _resolve_db(user, tenant_id)
        cursor = tdb["import_logs"].find({}).sort("created_at", -1).limit(50)
        logs = await cursor.to_list(length=50)
        for log in logs:
            log["id"] = str(log.pop("_id"))
        return logs

    @router.get("/template/{collection}")
    async def download_template(
        collection: str,
        format: str = Query("csv", enum=["csv", "xlsx"]),
        user: dict = Depends(get_current_user)
    ):
        """Download empty template for import"""
        if collection not in EXPORTABLE_COLLECTIONS:
            raise HTTPException(status_code=400, detail="Invalid collection")

        fields = EXPORTABLE_COLLECTIONS[collection]["fields"]

        if format == "csv":
            output = io.StringIO()
            writer = csv.writer(output)
            writer.writerow(fields)
            # Add one sample row
            writer.writerow(["" for _ in fields])
            output.seek(0)
            return StreamingResponse(
                iter([output.getvalue()]),
                media_type="text/csv",
                headers={"Content-Disposition": f'attachment; filename="{collection}_template.csv"'}
            )
        else:
            wb = Workbook()
            ws = wb.active
            ws.title = collection
            ws.append(fields)
            for col in ws.columns:
                ws.column_dimensions[col[0].column_letter].width = 20
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            return StreamingResponse(
                output,
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f'attachment; filename="{collection}_template.xlsx"'}
            )

    return router
